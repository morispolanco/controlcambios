import streamlit as st
from docx import Document
from googletrans import Translator

def translate_to_english(document):
    translator = Translator()
    document_text = ""
    for paragraph in document.paragraphs:
        document_text += paragraph.text + "\n"
    translated_text = translator.translate(document_text, src='es', dest='en').text
    
    document_en = Document()
    document_en.add_paragraph(translated_text)
    
    return document_en

def translate_to_spanish(document_en):
    translator = Translator()
    document_text_en = ""
    for paragraph in document_en.paragraphs:
        document_text_en += paragraph.text + "\n"
    translated_text_es = translator.translate(document_text_en, src='en', dest='es').text
    
    document_es = Document()
    document_es.add_paragraph(translated_text_es)
    
    return document_es

def generate_comparison_document(document, document_es):
    comparison_doc = Document()
    comparison_doc.add_heading('Comparison between the original and translated documents', level=1)
    
    for i in range(len(document.paragraphs)):
        p1 = document.paragraphs[i].text
        p2 = document_es.paragraphs[i].text
        
        if p1 != p2:
            run = comparison_doc.add_paragraph().add_run(p1)
            run.bold = True
            run.underline = True
            comparison_doc.add_paragraph(p2)
        else:
            comparison_doc.add_paragraph(p1)
    
    return comparison_doc

st.title("Document Translation App")

uploaded_file = st.file_uploader("Upload a .docx file", type="docx")

if uploaded_file is not None:
    document = Document(uploaded_file)
    
    st.header("Original Document (Spanish)")
    for paragraph in document.paragraphs:
        st.write(paragraph.text)
    
    st.header("Translated Document (English)")
    document_en = translate_to_english(document)
    for paragraph in document_en.paragraphs:
        st.write(paragraph.text)
    
    st.header("Translated Document (Back to Spanish)")
    document_es = translate_to_spanish(document_en)
    for paragraph in document_es.paragraphs:
        st.write(paragraph.text)
    
    st.header("Comparison Document with Track Changes")
    comparison_doc = generate_comparison_document(document, document_es)
    comparison_doc.save("comparison_document.docx")
    st.download_button(
        label="Download Comparison Document",
        data=comparison_doc.save,
        file_name="comparison_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
